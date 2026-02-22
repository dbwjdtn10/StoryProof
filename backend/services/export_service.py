"""
바이블 데이터 내보내기 서비스 (TXT / PDF / DOCX)
+ 챕터 본문 내보내기 서비스
"""

import io
import os
import re
from typing import Any, Dict, List

from fpdf import FPDF
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


FONTS_DIR = os.path.join(os.path.dirname(__file__), "..", "static", "fonts")


class BibleExportService:

    @staticmethod
    def filter_bible_data(bible_data: Dict[str, Any], search_query: str) -> Dict[str, Any]:
        """검색어로 바이블 데이터 필터링"""
        if not search_query or not search_query.strip():
            return bible_data

        q = search_query.strip().lower()
        result: Dict[str, Any] = {}

        # characters
        chars = bible_data.get("characters", [])
        result["characters"] = [
            c for c in chars
            if q in (c.get("name") or "").lower()
            or q in (c.get("description") or "").lower()
            or any(q in (t or "").lower() for t in (c.get("traits") or []))
        ]

        # items
        items = bible_data.get("items", [])
        result["items"] = [
            i for i in items
            if q in (i.get("name") or "").lower()
            or q in (i.get("description") or "").lower()
        ]

        # locations
        locations = bible_data.get("locations", [])
        result["locations"] = [
            loc for loc in locations
            if q in (loc.get("name") or "").lower()
            or q in (loc.get("description") or "").lower()
        ]

        # key_events
        events = bible_data.get("key_events", [])
        result["key_events"] = [
            e for e in events
            if q in (e.get("summary") or "").lower()
        ]

        # scenes
        scenes = bible_data.get("scenes", [])
        result["scenes"] = [
            s for s in scenes
            if q in (s.get("original_text") or "").lower()
            or q in (s.get("summary") or "").lower()
        ]

        # timeline (pass through)
        result["timeline"] = bible_data.get("timeline", [])

        return result

    @staticmethod
    def export_txt(bible_data: Dict[str, Any], title: str = "") -> bytes:
        """TXT 형식으로 바이블 데이터 내보내기"""
        lines: List[str] = []

        header = f"== 스토리보드 바이블 - {title} ==" if title else "== 스토리보드 바이블 =="
        lines.append(header)
        lines.append("=" * 60)
        lines.append("")

        # Characters
        chars = bible_data.get("characters", [])
        if chars:
            lines.append("[ 인물 ]")
            lines.append("")
            for c in chars:
                lines.append(f"  {c.get('name', '이름 없음')}")
                if c.get("description"):
                    lines.append(f"    설명: {c['description']}")
                if c.get("traits"):
                    lines.append(f"    특징: {', '.join(c['traits'])}")
                if c.get("appearance_count"):
                    lines.append(f"    등장: {c['appearance_count']}회")
                lines.append("")

        # Locations
        locations = bible_data.get("locations", [])
        if locations:
            lines.append("[ 장소 ]")
            lines.append("")
            for loc in locations:
                lines.append(f"  {loc.get('name', '이름 없음')}")
                if loc.get("description"):
                    lines.append(f"    설명: {loc['description']}")
                lines.append("")

        # Items
        items = bible_data.get("items", [])
        if items:
            lines.append("[ 아이템 ]")
            lines.append("")
            for item in items:
                lines.append(f"  {item.get('name', '이름 없음')}")
                if item.get("description"):
                    lines.append(f"    설명: {item['description']}")
                lines.append("")

        # Key Events
        events = bible_data.get("key_events", [])
        if events:
            lines.append("[ 주요 사건 ]")
            lines.append("")
            for e in events:
                scene_info = f" (Scene {e['scene_index'] + 1})" if e.get("scene_index") is not None else ""
                importance = f" [중요도: {e['importance']}]" if e.get("importance") else ""
                lines.append(f"  {e.get('summary', '')}{scene_info}{importance}")
            lines.append("")

        # Scenes
        scenes = bible_data.get("scenes", [])
        if scenes:
            lines.append("[ 씬 원문 ]")
            lines.append("")
            for s in scenes:
                idx = s.get("scene_index", 0)
                lines.append(f"--- Scene {idx + 1} ---")
                if s.get("summary"):
                    lines.append(f"요약: {s['summary']}")
                lines.append("")
                lines.append(s.get("original_text", ""))
                lines.append("")

        return "\n".join(lines).encode("utf-8")

    @staticmethod
    def export_pdf(bible_data: Dict[str, Any], title: str = "") -> bytes:
        """PDF 형식으로 바이블 데이터 내보내기 (한글 지원)"""
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)

        # Register Korean font
        font_regular = os.path.join(FONTS_DIR, "NanumGothic.ttf")
        font_bold = os.path.join(FONTS_DIR, "NanumGothicBold.ttf")

        pdf.add_font("NanumGothic", "", font_regular)
        if os.path.exists(font_bold):
            pdf.add_font("NanumGothic", "B", font_bold)

        pdf.add_page()

        # Title
        pdf.set_font("NanumGothic", "B", 16)
        header = f"스토리보드 바이블 - {title}" if title else "스토리보드 바이블"
        pdf.cell(0, 12, header, new_x="LMARGIN", new_y="NEXT", align="C")
        pdf.ln(8)

        def section_header(text: str):
            pdf.set_font("NanumGothic", "B", 13)
            pdf.set_fill_color(240, 240, 240)
            pdf.cell(0, 10, text, new_x="LMARGIN", new_y="NEXT", fill=True)
            pdf.ln(3)

        def body_text(text: str, indent: int = 0):
            pdf.set_font("NanumGothic", "", 10)
            x = pdf.l_margin + indent
            w = pdf.w - x - pdf.r_margin
            pdf.set_x(x)
            pdf.multi_cell(w, 6, text)

        def item_name(text: str):
            pdf.set_font("NanumGothic", "B", 11)
            pdf.cell(0, 8, text, new_x="LMARGIN", new_y="NEXT")

        # Characters
        chars = bible_data.get("characters", [])
        if chars:
            section_header("인물")
            for c in chars:
                item_name(c.get("name", "이름 없음"))
                if c.get("description"):
                    body_text(f"설명: {c['description']}", indent=5)
                if c.get("traits"):
                    body_text(f"특징: {', '.join(c['traits'])}", indent=5)
                if c.get("appearance_count"):
                    body_text(f"등장: {c['appearance_count']}회", indent=5)
                pdf.ln(3)

        # Locations
        locations = bible_data.get("locations", [])
        if locations:
            section_header("장소")
            for loc in locations:
                item_name(loc.get("name", "이름 없음"))
                if loc.get("description"):
                    body_text(f"설명: {loc['description']}", indent=5)
                pdf.ln(3)

        # Items
        items = bible_data.get("items", [])
        if items:
            section_header("아이템")
            for item in items:
                item_name(item.get("name", "이름 없음"))
                if item.get("description"):
                    body_text(f"설명: {item['description']}", indent=5)
                pdf.ln(3)

        # Key Events
        events = bible_data.get("key_events", [])
        if events:
            section_header("주요 사건")
            for e in events:
                scene_info = f" (Scene {e['scene_index'] + 1})" if e.get("scene_index") is not None else ""
                importance = f" [중요도: {e['importance']}]" if e.get("importance") else ""
                body_text(f"• {e.get('summary', '')}{scene_info}{importance}")
            pdf.ln(3)

        # Scenes
        content_w = pdf.w - pdf.l_margin - pdf.r_margin
        scenes = bible_data.get("scenes", [])
        if scenes:
            section_header("씬 원문")
            for s in scenes:
                idx = s.get("scene_index", 0)
                pdf.set_font("NanumGothic", "B", 10)
                pdf.cell(0, 8, f"--- Scene {idx + 1} ---", new_x="LMARGIN", new_y="NEXT")
                if s.get("summary"):
                    pdf.set_font("NanumGothic", "", 9)
                    pdf.set_text_color(100, 100, 100)
                    pdf.set_x(pdf.l_margin)
                    pdf.multi_cell(content_w, 5, f"요약: {s['summary']}")
                    pdf.set_text_color(0, 0, 0)
                body_text(s.get("original_text", ""))
                pdf.ln(5)

        return bytes(pdf.output())

    @staticmethod
    def export_docx(bible_data: Dict[str, Any], title: str = "") -> bytes:
        """DOCX 형식으로 바이블 데이터 내보내기 (한글 폰트 적용)"""
        doc = Document()

        # 기본 폰트를 한글 호환 폰트로 설정
        style = doc.styles["Normal"]
        font = style.font
        font.name = "맑은 고딕"
        font.size = Pt(11)

        # Title
        header = f"스토리보드 바이블 - {title}" if title else "스토리보드 바이블"
        title_para = doc.add_heading(header, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title_para.runs:
            run.font.name = "맑은 고딕"

        # Characters
        chars = bible_data.get("characters", [])
        if chars:
            doc.add_heading("인물", level=1)
            for c in chars:
                doc.add_heading(c.get("name", "이름 없음"), level=2)
                if c.get("description"):
                    doc.add_paragraph(f"설명: {c['description']}")
                if c.get("traits"):
                    doc.add_paragraph(f"특징: {', '.join(c['traits'])}")
                if c.get("appearance_count"):
                    doc.add_paragraph(f"등장: {c['appearance_count']}회")

        # Locations
        locations = bible_data.get("locations", [])
        if locations:
            doc.add_heading("장소", level=1)
            for loc in locations:
                doc.add_heading(loc.get("name", "이름 없음"), level=2)
                if loc.get("description"):
                    doc.add_paragraph(f"설명: {loc['description']}")

        # Items
        items = bible_data.get("items", [])
        if items:
            doc.add_heading("아이템", level=1)
            for item in items:
                doc.add_heading(item.get("name", "이름 없음"), level=2)
                if item.get("description"):
                    doc.add_paragraph(f"설명: {item['description']}")

        # Key Events
        events = bible_data.get("key_events", [])
        if events:
            doc.add_heading("주요 사건", level=1)
            for e in events:
                scene_info = f" (Scene {e['scene_index'] + 1})" if e.get("scene_index") is not None else ""
                importance = f" [중요도: {e['importance']}]" if e.get("importance") else ""
                doc.add_paragraph(
                    f"{e.get('summary', '')}{scene_info}{importance}",
                    style="List Bullet"
                )

        # Scenes
        scenes = bible_data.get("scenes", [])
        if scenes:
            doc.add_heading("씬 원문", level=1)
            for s in scenes:
                idx = s.get("scene_index", 0)
                doc.add_heading(f"Scene {idx + 1}", level=2)
                if s.get("summary"):
                    summary_para = doc.add_paragraph()
                    run = summary_para.add_run(f"요약: {s['summary']}")
                    run.italic = True
                    run.font.size = Pt(9)
                doc.add_paragraph(s.get("original_text", ""))

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()


class ChapterExportService:
    """챕터 본문 HTML을 TXT / PDF / DOCX로 변환"""

    @staticmethod
    def html_to_plain(html: str) -> str:
        """HTML을 plain text로 변환 (p/br → 줄바꿈, 태그 제거)"""
        if not html:
            return ""
        text = html
        # <br>, <br/>, <br /> → 줄바꿈
        text = re.sub(r"<br\s*/?>", "\n", text, flags=re.IGNORECASE)
        # </p>, </div>, </li> → 줄바꿈
        text = re.sub(r"</(?:p|div|li|h[1-6])>", "\n", text, flags=re.IGNORECASE)
        # <li> → bullet prefix
        text = re.sub(r"<li[^>]*>", "  - ", text, flags=re.IGNORECASE)
        # 나머지 태그 제거
        text = re.sub(r"<[^>]+>", "", text)
        # HTML entities
        text = text.replace("&nbsp;", " ").replace("&amp;", "&")
        text = text.replace("&lt;", "<").replace("&gt;", ">")
        text = text.replace("&quot;", '"')
        # 연속 빈줄 정리
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    @staticmethod
    def _parse_html_blocks(html: str) -> List[Dict[str, Any]]:
        """HTML을 블록 단위로 파싱 (heading, paragraph, list item 등)"""
        blocks: List[Dict[str, Any]] = []
        if not html:
            return blocks

        # heading 패턴
        heading_pat = re.compile(
            r"<(h[1-6])[^>]*>(.*?)</\1>", re.IGNORECASE | re.DOTALL
        )
        # paragraph/div 패턴
        para_pat = re.compile(
            r"<(p|div)[^>]*>(.*?)</\1>", re.IGNORECASE | re.DOTALL
        )
        # list item 패턴
        li_pat = re.compile(r"<li[^>]*>(.*?)</li>", re.IGNORECASE | re.DOTALL)

        # 순서대로 매칭하기 위해 통합 패턴 사용
        combined = re.compile(
            r"<(h[1-6]|p|div|li)[^>]*>(.*?)</\1>",
            re.IGNORECASE | re.DOTALL,
        )

        for m in combined.finditer(html):
            tag = m.group(1).lower()
            inner = m.group(2)

            if tag.startswith("h") and len(tag) == 2:
                level = int(tag[1])
                plain = re.sub(r"<[^>]+>", "", inner).strip()
                if plain:
                    blocks.append({"type": "heading", "level": level, "text": plain})
            elif tag == "li":
                plain = re.sub(r"<[^>]+>", "", inner).strip()
                if plain:
                    blocks.append({"type": "list_item", "text": plain})
            else:
                # p / div — inline 서식 유지
                blocks.append({"type": "paragraph", "html": inner})

        # 매칭이 없으면 전체를 plain text paragraph로
        if not blocks:
            plain = ChapterExportService.html_to_plain(html)
            if plain:
                blocks.append({"type": "paragraph", "html": plain})

        return blocks

    @staticmethod
    def _strip_tags(html: str) -> str:
        """태그 제거 + 기본 entity 변환"""
        text = re.sub(r"<br\s*/?>", "\n", html, flags=re.IGNORECASE)
        text = re.sub(r"<[^>]+>", "", text)
        text = text.replace("&nbsp;", " ").replace("&amp;", "&")
        text = text.replace("&lt;", "<").replace("&gt;", ">")
        text = text.replace("&quot;", '"')
        return text.strip()

    # ---- TXT ----

    @staticmethod
    def export_chapter_txt(content_html: str, title: str = "") -> bytes:
        """챕터 본문을 TXT로 내보내기"""
        lines: List[str] = []
        if title:
            lines.append(title)
            lines.append("=" * 40)
            lines.append("")

        plain = ChapterExportService.html_to_plain(content_html)
        lines.append(plain)
        return "\n".join(lines).encode("utf-8")

    # ---- PDF ----

    @staticmethod
    def export_chapter_pdf(content_html: str, title: str = "") -> bytes:
        """챕터 본문을 PDF로 내보내기 (NanumGothic 한글 폰트)"""
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)

        font_regular = os.path.join(FONTS_DIR, "NanumGothic.ttf")
        font_bold = os.path.join(FONTS_DIR, "NanumGothicBold.ttf")

        pdf.add_font("NanumGothic", "", font_regular)
        if os.path.exists(font_bold):
            pdf.add_font("NanumGothic", "B", font_bold)

        pdf.add_page()
        content_w = pdf.w - pdf.l_margin - pdf.r_margin

        # 제목
        if title:
            pdf.set_font("NanumGothic", "B", 16)
            pdf.cell(0, 12, title, new_x="LMARGIN", new_y="NEXT", align="C")
            pdf.ln(8)

        # html_to_plain으로 전체 텍스트 추출 (TXT와 동일 — 정상 동작 확인됨)
        plain = ChapterExportService.html_to_plain(content_html)
        paragraphs = plain.split("\n")

        pdf.set_font("NanumGothic", "", 10)
        for para in paragraphs:
            if para.strip():
                pdf.multi_cell(content_w, 6, para)
                pdf.ln(2)
            else:
                pdf.ln(4)

        return bytes(pdf.output())

    # ---- DOCX ----

    @staticmethod
    def export_chapter_docx(content_html: str, title: str = "") -> bytes:
        """챕터 본문을 DOCX로 내보내기 (한글 폰트 적용)"""
        doc = Document()

        # 기본 폰트를 한글 호환 폰트로 설정
        style = doc.styles["Normal"]
        font = style.font
        font.name = "맑은 고딕"
        font.size = Pt(11)

        if title:
            t = doc.add_heading(title, level=0)
            t.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in t.runs:
                run.font.name = "맑은 고딕"

        # html_to_plain으로 전체 텍스트 추출 (TXT와 동일 — 정상 동작 확인됨)
        plain = ChapterExportService.html_to_plain(content_html)
        paragraphs = plain.split("\n")

        for para_text in paragraphs:
            if para_text.strip():
                para = doc.add_paragraph()
                run = para.add_run(para_text)
                run.font.name = "맑은 고딕"
                run.font.size = Pt(11)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    @staticmethod
    def _add_formatted_runs(para, html: str):
        """HTML inline 서식(bold, italic)을 docx run으로 변환"""
        # <strong>/<b> → bold, <em>/<i> → italic, 나머지 → normal
        # 간단한 토큰 분리
        token_pat = re.compile(
            r"(</?(?:strong|b|em|i|u)>|<br\s*/?>)", re.IGNORECASE
        )
        parts = token_pat.split(html)

        bold = False
        italic = False
        underline = False

        for part in parts:
            lower = part.lower().strip()
            if lower in ("<strong>", "<b>"):
                bold = True
                continue
            elif lower in ("</strong>", "</b>"):
                bold = False
                continue
            elif lower in ("<em>", "<i>"):
                italic = True
                continue
            elif lower in ("</em>", "</i>"):
                italic = False
                continue
            elif lower in ("<u>",):
                underline = True
                continue
            elif lower in ("</u>",):
                underline = False
                continue
            elif lower.startswith("<br"):
                run = para.add_run("\n")
                continue

            # 나머지 태그 제거
            text = re.sub(r"<[^>]+>", "", part)
            text = text.replace("&nbsp;", " ").replace("&amp;", "&")
            text = text.replace("&lt;", "<").replace("&gt;", ">")
            text = text.replace("&quot;", '"')
            if not text:
                continue

            run = para.add_run(text)
            run.bold = bold
            run.italic = italic
            run.underline = underline
            run.font.size = Pt(11)
